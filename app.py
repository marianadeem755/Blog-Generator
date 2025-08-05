import streamlit as st
import os
import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import LETTER
from io import BytesIO
import markdown2
import re

# Load API Key
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama-3.3-70b-versatile"

# ---- Inject Custom CSS ----
st.markdown("""
    <style>
    .stApp {
        background-color: #121212 !important;
        font-family: 'Segoe UI', sans-serif;
        color: #ffffff;
    }
    section[data-testid="stSidebar"] > div:first-child {
        background-color: #1f1f1f;
        padding: 30px 10px 20px 10px;
    }
    .sidebar-links a {
        font-size: 18px;
        color: #ffcc00;
        text-decoration: none;
    }
    .sidebar-links a:hover {
        color: #00ffff;
        text-decoration: underline;
    }
    h1, h3, h2, h4 {
        color: #ffd700;
        text-shadow: 1px 1px 3px rgba(0,0,0,0.4);
    }
    p {
        color: #f1f1f1;
    }
    .stButton button, .download-button button {
        background: linear-gradient(90deg, #f7971e, #ffd200) !important;
        color: #000000 !important;
        font-weight: bold;
        font-size: 18px !important;
        border-radius: 12px;
        padding: 10px 20px;
        box-shadow: 0 4px 10px rgba(0, 0, 0, 0.3);
    }
    .stButton button:hover, .download-button button:hover {
        transform: scale(1.02);
        box-shadow: 0 6px 15px rgba(0, 0, 0, 0.4);
    }
    .connect-title {
        font-size: 24px;
        color: #32a852;
        font-weight: bold;
        margin-bottom: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# ---- Sidebar ----
st.sidebar.markdown("""<div class="connect-title">🔗 Connect With Me</div>""", unsafe_allow_html=True)
st.sidebar.markdown("""<div class="sidebar-links">
<a href="https://github.com/marianadeem755" target="_blank">🌐 GitHub</a><br>
<a href="https://www.kaggle.com/marianadeem755" target="_blank">📊 Kaggle</a><br>
<a href="mailto:marianadeem755@gmail.com">📧 Email</a><br>
<a href="https://huggingface.co/maria355" target="_blank">🤗 Hugging Face</a>
</div>""", unsafe_allow_html=True)

st.sidebar.markdown("---")
st.sidebar.markdown("""<span style='color: #32a852; font-size: 18px;'>📖 About This App</span>""", unsafe_allow_html=True)
st.sidebar.markdown("""
This app uses the powerful LLaMA 3 AI model to generate high-quality blog posts in minutes. 
Simply enter a blog topic, select the tone, and choose your preferred download format (Markdown, PDF, or Word).
""", unsafe_allow_html=True)

# ---- Title and Inputs ----
st.markdown("<h1>🌟 AI Blog Generator</h1>", unsafe_allow_html=True)
st.markdown("<p style='font-size: 20px;'>Create blog posts in minutes using LLaMA 3 model.</p>", unsafe_allow_html=True)
st.markdown("---")

col1, col2 = st.columns(2)
with col1:
    topic = st.text_input("📌 Enter Blog Topic", placeholder="e.g., Future of AI in Education")
with col2:
    tone = st.selectbox("🎙️ Select Tone", ["Informative", "Conversational", "Professional", "Casual"])

format_choice = st.radio("💾 Choose Download Format", ["Markdown", "PDF", "Word (.docx)"], horizontal=True)

# ---- Generate Button and Logic ----
st.markdown("<br>", unsafe_allow_html=True)
if st.button("🚀 Generate Blog Post", use_container_width=True):
    if not topic:
        st.warning("Please enter a blog topic.", icon="⚠️")
    elif not GROQ_API_KEY:
        st.error("GROQ_API_KEY not set. Please check your .env file.", icon="❌")
    else:
        with st.spinner("Crafting your blog with AI magic... ✨"):
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": GROQ_MODEL,
                "messages": [
                    {"role": "system", "content": "You are a professional blog writer and researcher."},
                    {"role": "user", "content": f"""
Write a high-quality, well-structured blog post on the topic: "{topic}". The blog post should:

1. Be engaging, informative, and human-like in tone ({tone}).
2. Use Markdown-style headings and subheadings.
3. Be written in simple, reader-friendly language.
4. Be free of plagiarism and fluff.
5. Include facts, examples, and if appropriate, recent statistics or case studies.

Start writing the blog now.
"""}
                ],
                "temperature": 0.7,
                "max_tokens": 1500
            }

            try:
                response = requests.post(GROQ_API_URL, headers=headers, json=payload)
                response.raise_for_status()
                blog = response.json()["choices"][0]["message"]["content"]

                st.markdown("### ✨ Generated Blog Post")
                st.markdown(blog)

                if format_choice == "Markdown":
                    st.download_button("📥 Download Markdown", blog, file_name="blog_post.md", key="md", type="primary", help="Download as Markdown")

                elif format_choice == "PDF":
                    html = markdown2.markdown(blog)
                    pdf_buffer = BytesIO()
                    doc = SimpleDocTemplate(pdf_buffer, pagesize=LETTER)
                    styles = getSampleStyleSheet()
                    story = []
                    for paragraph in html.split("\n"):
                        if paragraph.strip():
                            story.append(Paragraph(paragraph, styles["Normal"]))
                            story.append(Spacer(1, 12))
                    doc.build(story)
                    pdf_buffer.seek(0)
                    st.download_button("📥 Download PDF", pdf_buffer, file_name="blog_post.pdf", key="pdf", type="primary", help="Download as PDF")

                elif format_choice == "Word (.docx)":
                    doc = Document()
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Segoe UI'
                    font.size = Pt(11)

                    for line in blog.split('\n'):
                        line = line.strip()
                        if not line or line in ['---', '===', '---', '___']:
                            continue  # Skip Markdown separators

                        # Clean markdown formatting for Word
                        line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # remove bold
                        line = re.sub(r'\*(.*?)\*', r'\1', line)      # remove italic
                        line = re.sub(r'^-{3,}$', '', line)            # remove --- lines
                        line = re.sub(r'^={3,}$', '', line)            # remove === lines

                        if re.match(r"^#{1,6} ", line):
                            level = line.count("#", 0, line.find(" "))
                            heading = line.replace("#", "").strip()
                            doc.add_heading(heading, level=min(level, 4))
                        elif re.match(r"^[-*] ", line):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        else:
                            doc.add_paragraph(line)

                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    st.download_button("📥 Download Word", doc_buffer, file_name="blog_post.docx", key="docx", type="primary", help="Download as DOCX")

            except Exception as e:
                st.error(f"Failed to generate blog post: {e}", icon="❌")

